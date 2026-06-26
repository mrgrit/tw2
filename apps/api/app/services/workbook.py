"""학생 워크북(.docx) 렌더링 — 빈 워크북(사전 배포)과 **채워진 워크북(사후 복습)** 단일 생성기.

핵심: 미션마다 3개의 붙여넣기 칸(▶실행 명령/페이로드 · ▶실행 결과 · ▶설명/분석)이 곧 학생
제출 저널(StudentSubmission)의 필드(what_i_did / what_happened / description)와 1:1 대응한다.
`prefill` 을 주면 그 칸을 학생 입력으로 자동 채우고, 채점 결과(verdict/점수/피드백)도 함께 싣는다.
주지 않으면 기존처럼 빈 칸(사전 배포용)으로 남는다.

- `build_doc(scn, prefill=None) -> Document` : scn 은 워크북 dict(아래 scenario_to_dict 참고).
- `scenario_to_dict(scenario)` : DB Scenario(ORM) → 워크북 dict 어댑터.
- `prefill_from_submissions(subs)` : 제출 저널 → {(side,order): {...}} (미션별 최신 1건).

scripts/gen_workbooks.py(YAML 일괄)와 API 다운로드(DB+제출)가 같은 렌더러를 공유한다.
"""
from __future__ import annotations
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "맑은 고딕"
MONO = "Consolas"


def _set_cell_border(cell, color="999999", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def _shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _kfont(run):
    run.font.name = KFONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), KFONT)


def paste_box(doc, height_cm, hint="", content="", mono=False):
    """테두리 칸. content 가 있으면 학생 입력으로 채우고, 없으면 회색 안내(hint)만 둔다."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _set_cell_border(cell)
    row = t.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p = cell.paragraphs[0]
    if content:
        for j, line in enumerate(str(content).replace("\r\n", "\n").split("\n")):
            r = p.add_run(line)
            if mono:
                r.font.name = MONO
                r.font.size = Pt(9.5)
            else:
                r.font.size = Pt(10)
                _kfont(r)
            if j < len(str(content).split("\n")) - 1:
                p.add_run().add_break()
    elif hint:
        r = p.add_run(hint)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        _kfont(r)
    return t


def add_inline(p, text, size=10.5):
    pos = 0
    for m in re.finditer(r"\*\*([^*]+)\*\*|`([^`]+)`", text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.font.size = Pt(size)
            _kfont(r)
        if m.group(1) is not None:
            r = p.add_run(m.group(1))
            r.bold = True
            r.font.size = Pt(size)
            _kfont(r)
        else:
            r = p.add_run(m.group(2))
            r.font.name = MONO
            r.font.size = Pt(size - 0.5)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.font.size = Pt(size)
        _kfont(r)


def render_md(doc, text):
    lines = (text or "").replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if s.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            _shade(p)
            r = p.add_run("\n".join(code))
            r.font.name = MONO
            r.font.size = Pt(9.5)
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            lvl = len(h.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(h.group(2))
            r.bold = True
            r.font.size = Pt(13 if lvl <= 2 else 11.5)
            _kfont(r)
            i += 1
            continue
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            content = re.sub(r"^\s*([-*]|\d+\.)\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue
        if s == "":
            i += 1
            continue
        para = []
        while (i < len(lines) and lines[i].strip() != ""
               and not lines[i].strip().startswith("```")
               and not re.match(r"^(#{1,6})\s+", lines[i])
               and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i])):
            para.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        for j, pl in enumerate(para):
            add_inline(p, pl.strip())
            if j < len(para) - 1:
                p.add_run().add_break()


_VERDICT_KO = {"pass": "통과", "partial": "부분", "fail": "미통과", "review": "검토대기"}


def _grade_block(doc, pf: dict):
    """채워진 워크북 한정 — 미션 칸 아래 'AI 채점 결과'(verdict/점수/피드백/기준)."""
    status = pf.get("grade_status")
    if not status or status == "pending":
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("🧑‍🏫 AI 채점 결과")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x6F, 0x3C)
    _kfont(r)
    if status == "failed":
        fr = doc.add_paragraph().add_run("채점 보류(강사 검토 대상)")
        fr.font.size = Pt(10); _kfont(fr)
        return
    v = pf.get("verdict")
    aw, mx = pf.get("awarded_points"), pf.get("max_points")
    head = doc.add_paragraph()
    hr = head.add_run(f"판정: {_VERDICT_KO.get(v, v or '-')}    점수: {aw if aw is not None else '-'}"
                      f"{('/' + str(mx)) if mx is not None else ''}")
    hr.font.size = Pt(10.5); hr.bold = True; _kfont(hr)
    if pf.get("feedback"):
        render_md(doc, pf["feedback"])
    for label, items in (("충족", pf.get("criteria_met")), ("미충족", pf.get("criteria_missing"))):
        for it in (items or []):
            bp = doc.add_paragraph(style="List Bullet")
            br = bp.add_run(f"[{label}] {it}")
            br.font.size = Pt(9.5); _kfont(br)


def build_doc(scn: dict, prefill: dict | None = None) -> Document:
    """워크북 1부 생성. prefill={(side,order): {what_i_did,...,verdict,...}} 면 채워진 워크북."""
    pf_map = prefill or {}
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = KFONT
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KFONT)

    title = doc.add_heading(level=0)
    tr = title.add_run(scn.get("title", scn.get("id", "")))
    _kfont(tr)

    meta = doc.add_paragraph()
    mins = (scn.get("time_limit") or 0) // 60
    mr = meta.add_run(f"카테고리 {scn.get('category','-')}  ·  난이도 {scn.get('difficulty','-')}"
                      f"  ·  유형 {scn.get('battle_type','-')}  ·  제한 {mins}분")
    mr.italic = True
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _kfont(mr)

    info = doc.add_paragraph()
    nm = scn.get("student_name") or "__________________"
    ir = info.add_run(f"이름: {nm}    학번: __________________    날짜: __________________")
    ir.font.size = Pt(11)
    _kfont(ir)

    doc.add_paragraph()
    ov = doc.add_paragraph()
    ovr = ov.add_run("■ 시나리오 개요")
    ovr.bold = True
    ovr.font.size = Pt(13)
    _kfont(ovr)
    render_md(doc, scn.get("description", ""))

    reds = scn.get("red_missions") or []
    blues = scn.get("blue_missions") or []
    sections = [("🔴 RED", "red", m) for m in reds] + [("🔵 BLUE", "blue", m) for m in blues]
    for label, side, m in sections:
        doc.add_page_break()
        order = m.get("order", "")
        head = doc.add_paragraph()
        hr = head.add_run(f"{label} #{order}   (배점 {m.get('points',0)}점  ·  대상 VM: {m.get('target_vm','-')})")
        hr.bold = True
        hr.font.size = Pt(15)
        hr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if side == "red" else RGBColor(0x00, 0x40, 0xA0)
        _kfont(hr)

        render_md(doc, m.get("instruction", ""))

        pf = pf_map.get((side, order)) or pf_map.get((side, str(order))) or {}
        doc.add_paragraph()
        boxes = [
            ("▶ 실행한 명령 / 페이로드", 3.0, "여기에 실행한 명령·페이로드를 입력하거나 캡처를 붙여넣으세요",
             pf.get("what_i_did", ""), True),
            ("▶ 실행 결과 (출력/스크린샷)", 8.5, "여기에 결과 출력을 적거나 스크린샷을 붙여넣으세요 (커서를 두고 Ctrl+V)",
             pf.get("what_happened", ""), True),
            ("▶ 설명 / 분석 (무엇을·왜 했고 결과를 어떻게 해석했나)", 3.5, "",
             pf.get("description", ""), False),
        ]
        for lab, h, hint, content, mono in boxes:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(6)
            lr = lp.add_run(lab)
            lr.bold = True
            lr.font.size = Pt(11)
            _kfont(lr)
            paste_box(doc, h, hint=hint, content=content, mono=mono)

        if pf:
            _grade_block(doc, pf)

    return doc


# ── training 실습(lab) → 학생 워크북 ──
def build_lab_doc(lab: dict, track_label: str = "", week: int | None = None) -> Document:
    """training lab(yaml dict) → 학생 워크북 1부(따라하기 형식).

    각 step = instruction(목표·명령) + ✅이렇게 나오면 정상(expected_output) + 💡결과 해석
    (answer_detail) 참고 + 붙여넣기 3칸(학생이 직접 실행한 명령/결과/분석 기록).
    배틀 워크북과 동일 렌더러(render_md/paste_box) 공유. 명령은 instruction 의 💻 블록에 있으므로
    answer 를 따로 싣지 않고, 숨어 있던 정상 출력·해석을 표면화해 '시험지'가 아닌 '따라하기'로 만든다."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = KFONT
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KFONT)

    title = doc.add_heading(level=0)
    wk = f"W{week:02d} " if week else ""
    tr = title.add_run(f"{wk}{lab.get('title', '실습')}")
    _kfont(tr)

    meta = doc.add_paragraph()
    mins = lab.get("duration_minutes", 0)
    mr = meta.add_run(f"{track_label or lab.get('course','-')}  ·  난이도 {lab.get('difficulty','-')}"
                      f"  ·  제한 {mins}분  ·  합격선 {int((lab.get('pass_threshold',0) or 0)*100)}%")
    mr.italic = True
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _kfont(mr)

    info = doc.add_paragraph()
    ir = info.add_run("이름: __________________    학번: __________________    날짜: __________________")
    ir.font.size = Pt(11)
    _kfont(ir)

    doc.add_paragraph()
    if lab.get("description"):
        ov = doc.add_paragraph()
        ovr = ov.add_run("■ 실습 개요")
        ovr.bold = True; ovr.font.size = Pt(13); _kfont(ovr)
        render_md(doc, lab["description"])
    if lab.get("objectives"):
        ob = doc.add_paragraph()
        obr = ob.add_run("■ 학습 목표")
        obr.bold = True; obr.font.size = Pt(13); _kfont(obr)
        for o in lab["objectives"]:
            bp = doc.add_paragraph(style="List Bullet")
            add_inline(bp, str(o))

    for st in sorted(lab.get("steps", []), key=lambda x: x.get("order", 0)):
        doc.add_page_break()
        order = st.get("order", "")
        head = doc.add_paragraph()
        hr = head.add_run(f"🧪 STEP {order}   (배점 {st.get('points',0)}점  ·  대상 {st.get('target_vm','-')}"
                          f"  ·  {st.get('category','-')})")
        hr.bold = True
        hr.font.size = Pt(15)
        hr.font.color.rgb = RGBColor(0x00, 0x40, 0xA0)
        _kfont(hr)

        render_md(doc, st.get("instruction", ""))

        # ── 따라하기 참고: ✅ 정상 출력(expected_output) + 💡 결과 해석(answer_detail) 표면화 ──
        eo = str(st.get("expected_output") or "").strip()
        if eo:
            lp = doc.add_paragraph(); lp.paragraph_format.space_before = Pt(6)
            lr = lp.add_run("✅ 이렇게 나오면 정상 (실측 예 — 숫자·시간은 환경마다 다름)")
            lr.bold = True; lr.font.size = Pt(11); lr.font.color.rgb = RGBColor(0x1F, 0x6F, 0x3C); _kfont(lr)
            cp = doc.add_paragraph(); _shade(cp, "EAF5EA")
            cr = cp.add_run(eo); cr.font.name = MONO; cr.font.size = Pt(9.5)
        ad = str(st.get("answer_detail") or "").strip()
        if ad:
            lp = doc.add_paragraph(); lp.paragraph_format.space_before = Pt(6)
            lr = lp.add_run("💡 결과 해석")
            lr.bold = True; lr.font.size = Pt(11); lr.font.color.rgb = RGBColor(0x1F, 0x6F, 0x3C); _kfont(lr)
            render_md(doc, ad)

        doc.add_paragraph()
        boxes = [
            ("▶ 실행한 명령 / 페이로드", 3.0, "여기에 실행한 명령을 입력하거나 캡처를 붙여넣으세요", True),
            ("▶ 실행 결과 (출력/스크린샷)", 8.5, "여기에 결과 출력을 적거나 스크린샷을 붙여넣으세요 (Ctrl+V)", True),
            ("▶ 설명 / 분석 (무엇을·왜 했고 결과를 어떻게 해석했나)", 3.5, "", False),
        ]
        for lab_, h, hint, mono in boxes:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(6)
            lr = lp.add_run(lab_)
            lr.bold = True; lr.font.size = Pt(11); _kfont(lr)
            paste_box(doc, h, hint=hint, content="", mono=mono)

    return doc


# ── 어댑터: DB ORM / 제출 저널 → 렌더러 입력 ──
def scenario_to_dict(scenario, student_name: str | None = None) -> dict:
    """DB Scenario(ORM) → 워크북 dict. mission_red/blue 의 missions 배열을 펼친다."""
    def missions(side_json):
        return list((side_json or {}).get("missions") or [])
    mr, mb = scenario.mission_red or {}, scenario.mission_blue or {}
    return {
        "id": scenario.id,
        "title": scenario.title,
        "category": scenario.category or "-",
        "difficulty": (scenario.scoring or {}).get("difficulty", "-"),
        "battle_type": mr.get("battle_type") or mb.get("battle_type") or "-",
        "time_limit": scenario.time_limit_sec,
        "description": scenario.description or "",
        "red_missions": missions(mr),
        "blue_missions": missions(mb),
        "student_name": student_name,
    }


def dict_from_snapshots(subs: list, title: str | None = None,
                        student_name: str | None = None) -> dict:
    """시나리오가 삭제돼도 — 제출 저널의 mission_snapshot 들로 워크북 골격을 복원한다."""
    reds, blues, seen = [], [], set()
    for s in sorted(subs, key=lambda x: x.id):
        snap = s.mission_snapshot or {}
        side, order = s.mission_side, s.mission_order
        if not side or order is None or (side, order) in seen:
            continue
        seen.add((side, order))
        m = {"order": order, "points": snap.get("points", 0),
             "target_vm": snap.get("target_vm", "-"), "instruction": snap.get("instruction", "")}
        (reds if side == "red" else blues).append(m)
    reds.sort(key=lambda m: m["order"]); blues.sort(key=lambda m: m["order"])
    return {
        "id": None, "title": title or "워크북(시나리오 기록)", "category": "-", "difficulty": "-",
        "battle_type": "-", "time_limit": 0, "description": "",
        "red_missions": reds, "blue_missions": blues, "student_name": student_name,
    }


def prefill_from_submissions(subs: list) -> dict:
    """제출 저널 → {(side, order): {입력+채점}} — 미션별 최신 1건(id 오름차순, 마지막 우선)."""
    out: dict = {}
    for s in sorted(subs, key=lambda x: x.id):
        if not s.mission_side or s.mission_order is None:
            continue
        out[(s.mission_side, s.mission_order)] = {
            "what_i_did": s.what_i_did, "what_happened": s.what_happened,
            "description": s.description, "grade_status": s.grade_status,
            "verdict": s.verdict, "awarded_points": s.awarded_points,
            "max_points": s.max_points, "feedback": s.feedback,
            "criteria_met": s.criteria_met, "criteria_missing": s.criteria_missing,
        }
    return out
