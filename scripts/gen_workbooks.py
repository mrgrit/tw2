#!/usr/bin/env python3
"""공방전 학생 워크북(.docx) 생성기.

각 시나리오 YAML(contents/battle-scenarios/*.yaml) → 워크북 1개:
  - 시나리오 개요(설명)
  - 미션별(RED/BLUE) 안내(상황/할 일/채점/합격기준) + 실행/결과/분석 붙여넣기 빈칸
출력: contents/battle-workbook/<id>.docx

사용: .venv/bin/python scripts/gen_workbooks.py
"""
import glob
import os
import re
import sys

import yaml
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "contents", "battle-scenarios")
OUT = os.path.join(ROOT, "contents", "battle-workbook")
KFONT = "맑은 고딕"
MONO = "Consolas"


def _set_cell_border(cell, color="999999", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def _shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _kfont(run):
    run.font.name = KFONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), KFONT)


def paste_box(doc, height_cm, hint=""):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _set_cell_border(cell)
    row = t.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p = cell.paragraphs[0]
    if hint:
        r = p.add_run(hint)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        _kfont(r)
    return t


def add_inline(p, text, size=10.5):
    pos = 0
    for m in re.finditer(r"\*\*([^*]+)\*\*|`([^`]+)`", text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.font.size = Pt(size)
            _kfont(r)
        if m.group(1) is not None:
            r = p.add_run(m.group(1))
            r.bold = True
            r.font.size = Pt(size)
            _kfont(r)
        else:
            r = p.add_run(m.group(2))
            r.font.name = MONO
            r.font.size = Pt(size - 0.5)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.font.size = Pt(size)
        _kfont(r)


def render_md(doc, text):
    lines = (text or "").replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if s.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            _shade(p)
            r = p.add_run("\n".join(code))
            r.font.name = MONO
            r.font.size = Pt(9.5)
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            lvl = len(h.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(h.group(2))
            r.bold = True
            r.font.size = Pt(13 if lvl <= 2 else 11.5)
            _kfont(r)
            i += 1
            continue
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            content = re.sub(r"^\s*([-*]|\d+\.)\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue
        if s == "":
            i += 1
            continue
        para = []
        while (i < len(lines) and lines[i].strip() != ""
               and not lines[i].strip().startswith("```")
               and not re.match(r"^(#{1,6})\s+", lines[i])
               and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i])):
            para.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        for j, pl in enumerate(para):
            add_inline(p, pl.strip())
            if j < len(para) - 1:
                p.add_run().add_break()


def build(scn, outpath):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = KFONT
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KFONT)

    title = doc.add_heading(level=0)
    tr = title.add_run(scn.get("title", scn.get("id", "")))
    _kfont(tr)

    meta = doc.add_paragraph()
    mins = (scn.get("time_limit") or 0) // 60
    mr = meta.add_run(f"카테고리 {scn.get('category','-')}  ·  난이도 {scn.get('difficulty','-')}"
                      f"  ·  유형 {scn.get('battle_type','-')}  ·  제한 {mins}분")
    mr.italic = True
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _kfont(mr)

    info = doc.add_paragraph()
    ir = info.add_run("이름: __________________    학번: __________________    날짜: __________________")
    ir.font.size = Pt(11)
    _kfont(ir)

    doc.add_paragraph()
    ov = doc.add_paragraph()
    ovr = ov.add_run("■ 시나리오 개요")
    ovr.bold = True
    ovr.font.size = Pt(13)
    _kfont(ovr)
    render_md(doc, scn.get("description", ""))

    reds = scn.get("red_missions") or []
    blues = scn.get("blue_missions") or []
    sections = [("🔴 RED", m) for m in reds] + [("🔵 BLUE", m) for m in blues]
    for label, m in sections:
        doc.add_page_break()
        head = doc.add_paragraph()
        hr = head.add_run(f"{label} #{m.get('order','')}   (배점 {m.get('points',0)}점  ·  대상 VM: {m.get('target_vm','-')})")
        hr.bold = True
        hr.font.size = Pt(15)
        hr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if "RED" in label else RGBColor(0x00, 0x40, 0xA0)
        _kfont(hr)

        render_md(doc, m.get("instruction", ""))

        doc.add_paragraph()
        for lab, h, hint in [
            ("▶ 실행한 명령 / 페이로드", 3.0, "여기에 실행한 명령·페이로드를 입력하거나 캡처를 붙여넣으세요"),
            ("▶ 실행 결과 화면 (스크린샷 붙여넣기)", 8.5, "여기에 결과 스크린샷을 붙여넣으세요 (커서를 두고 Ctrl+V)"),
            ("▶ 설명 / 분석 (무엇을·왜 했고 결과를 어떻게 해석했나)", 3.5, ""),
        ]:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(6)
            lr = lp.add_run(lab)
            lr.bold = True
            lr.font.size = Pt(11)
            _kfont(lr)
            paste_box(doc, h, hint)

    doc.save(outpath)


def main():
    os.makedirs(OUT, exist_ok=True)
    files = sorted(glob.glob(os.path.join(SRC, "*.yaml")))
    n = 0
    for f in files:
        d = yaml.safe_load(open(f, encoding="utf-8"))
        if not isinstance(d, dict) or not (d.get("red_missions") or d.get("blue_missions")):
            continue
        sid = d.get("id") or os.path.splitext(os.path.basename(f))[0]
        out = os.path.join(OUT, f"{sid}.docx")
        build(d, out)
        n += 1
        print(f"  ✓ {sid}.docx  (RED {len(d.get('red_missions') or [])} / BLUE {len(d.get('blue_missions') or [])})")
    print(f"\n총 {n}개 워크북 생성 → {OUT}")


if __name__ == "__main__":
    main()
